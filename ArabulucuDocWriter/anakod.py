from tkinter import*
from Vericekici import Vericekici
import tkinter as tk
from tkinter import ttk
import docx
LARGE_FONT= ("Verdana", 12)
NORM_FONT = ("Helvetica", 10)
SMALL_FONT = ("Helvetica", 8)
class Arabirim():
    def __init__(self):
        self.root=Tk()
        self.root.title("Arabulucu Döküman Hazırlayıcı")
        self.root.geometry("1350x800")

        #def  myClick():
            #myLabel = Label(self.root, text="Hello "+e.get())
            #myLabel.grid(row=1,column=1, columnspan=1, padx=10, pady=10)

        #Temel ve Başvuran Labelları
        büroLabel =Label(self.root, text="Arabuluculuk Bürosu :")
        başvuruNoLabel =Label(self.root, text="Başvuru Numarası :")
        basvurutarihilabel= Label(self.root, text="Başvuru Tarihi :")
        başvuranLabel=Label(self.root, text="Başvuran Bilgileri")
        başvuranTCLabel=Label(self.root, text="TC Numarası :")
        başvuranİsimLabel=Label(self.root, text="İsim ve Soyadı :")
        başvuranAdresLabel=Label(self.root, text="Adresi :")
        başvuranVekilLabel=Label(self.root, text="Vekili :")

        #Başvurulan Labelları



        #Labelların yerleştirilmesi

        büroLabel.grid(row=2, column=1,columnspan=2)
        başvuruNoLabel.grid(row=3, column=1,columnspan=2)
        basvurutarihilabel.grid(row=4, column=1,columnspan=2)
        başvuranLabel.grid(row=5, column=1,columnspan=2)
        başvuranTCLabel.grid(row=6, column=1,columnspan=2)
        başvuranİsimLabel.grid(row=7, column=1,columnspan=2)
        başvuranAdresLabel.grid(row=8, column=1,columnspan=2)
        başvuranVekilLabel.grid(row=9, column=1,columnspan=2)



    
        #copy paste alanı
        entryText = StringVar()
        self.text1=Text(self.root, height=35, width=80)
        self.text1.grid(column=40, columnspan=80,row=0,rowspan=60)

        #e.grid(column=32, columnspan=28, row=2, rowspan=58)

        #Temel Entryler
        self.buro_entry = StringVar()
        büroEntry=Entry(self.root, width=50, borderwidth=1, textvariable=self.buro_entry)
        self.basvuru_entry = StringVar()
        başvuruEntry =Entry(self.root, width=50, borderwidth=1, textvariable=self.basvuru_entry)
        self.basvurutarihi_entry = StringVar()
        basvurutarihiEntry= Entry(self.root, width=50, borderwidth=1, textvariable=self.basvurutarihi_entry)
        self.tc_entry = StringVar()
        başvuranTCEntry=Entry(self.root, width=50, borderwidth=1, textvariable=self.tc_entry)

        self.basvuran_isim_entry = StringVar()
        başvuranİsimEntry=Entry(self.root, width=50, borderwidth=1,textvariable=self.basvuran_isim_entry)
        self.basvuran_adres_entry=StringVar()
        başvuranAdresEntry=Entry(self.root, width=50, borderwidth=1,textvariable=self.basvuran_adres_entry)
        self.basvuran_vekil_entry=StringVar()
        başvuranVekilEntry=Entry(self.root, width=50, borderwidth=1,textvariable=self.basvuran_vekil_entry)

        #Başvurulan Entryleri


        


        #bilgi girme alanlarının arabirime konumu
        büroEntry.grid(row=2,column=3, columnspan=20)
        başvuruEntry.grid(row=3,column=3, columnspan=20)
        basvurutarihiEntry.grid(row=4,column=3, columnspan=20)
        başvuranTCEntry.grid(row=6,column=3, columnspan=20)
        başvuranİsimEntry.grid(row=7,column=3, columnspan=20)
        başvuranAdresEntry.grid(row=8,column=3, columnspan=20)
        başvuranVekilEntry.grid(row=9,column=3, columnspan=20)





        #komutları burada tanımlayacağım
        isimsatırlistesi=list()
            #self.text1.insert(END, text)

                #if line.startswith("TC"):
                    #büroEntry.insert(0,"öyleyse dans")
                    #print("öyleyse dans")
                    #adressatırlistesi.append(int(satırsayısı))
                    #print(adressatırlistesi[0:"end"])
                    #büroEntry.insert(0,adressatırlistesi[0:6])

            #for line in input:


        #düğmeleri tanımlama

        Button1 = Button(self.root, text="Veri al", command=lambda:self.verial())
        
        Button3 = Button(self.root, text="Verileri temizle", command=lambda:self.veritemizle())
        Button2 = Button(self.root, text="Veri yaz", command=lambda:self.veriyaz())
        Button4 = Button(self.root, text="Ticari Dava Şartı Arabuluculuk", command=lambda:self.veriyaz())
        Button5 = Button(self.root, text="Ticari Dava ilk Tutanak ", command=lambda:self.ticaridavailktutanak())
        Button6 = Button(self.root, text="Ticari Dava Son Tutanak", command=lambda:self.ticaridavasontutanak())
        Button7 = Button(self.root, text="Ticari Dava Şartı Arabuluculuk Davet Mektubu", command=lambda:self.ticarialabuluculukdavetmektubu())
        Button8 = Button(self.root, text="İş hukuku bilgilendirme Tutanağı", command=lambda:self.ishukukubilgilendirmetutanagi())
        Button9 = Button(self.root, text="Arabulucu Belirleme Tutanağı", command=lambda:self.arabuluculukbelirlemetutanagi())
        #düğmeleri yerleştirme
        Button1.grid(row=70,column=50)
        Button3.grid(row=70,column=60)
        Button2.grid(row=80,column=40)
        Button4.grid(row=80,column=50)
        Button5.grid(row=80,column=60)
        Button6.grid(row=90,column=40)
        Button7.grid(row=90,column=50)  
        Button8.grid(row=90,column=60)
        Button9.grid(row=90,column=40)     
        #Button1.grid(row=2,column=1, columnspan=1, padx=10, pady=10)
        self.root.mainloop()
    def verial(self):
        line_list = self.text1.get('1.0', 'end').split('\n')
            
        
        self.tc,self.isimsoyisim,self.adres,self.basvuranisimsoyisim,self.basvuranadres,self.hukukburo,self.basvurunumarasi,self.basvurutarihi,self.vekili,self.isimlistesi,self.adreslistesi,self.vekillistesi,self.verginolistesi,self.telnolistesi,self.tcnolistesi=Vericekici().vericek(line_list)
        self.vekillistem=[]
        self.adreslistem=[]
        for i in self.vekillistesi:
            for j in i:
                self.vekillistem.append(j)
        for i in self.adreslistesi:
            for j in i:
                self.adreslistem.append(j)
        self.vekillistesi=self.vekillistem
        self.adreslistesi=self.adreslistem

        self.tc_entry.set(self.tc)
        self.basvuran_isim_entry.set(self.isimsoyisim)
        self.basvuran_adres_entry.set(self.adres)
        self.basvuran_vekil_entry.set(self.vekili)
        self.buro_entry.set(self.hukukburo)
        self.basvuru_entry.set(self.basvurunumarasi)
        self.basvurutarihi_entry.set(self.basvurutarihi)


        self.islist=list(range(0,len(self.isimlistesi)))
        self.adreslist=list(range(0,len(self.isimlistesi)))
        self.vekillist=list(range(0,len(self.isimlistesi)))
        self.tclist=list(range(0,len(self.isimlistesi)))
        self.verginolist=list(range(0,len(self.isimlistesi)))
        self.telnolist=list(range(0,len(self.isimlistesi)))
        
        row=11
        for i in range(0,len(self.isimlistesi)):
            self.islist[i] = StringVar()
            self.adreslist[i]= StringVar()
            self.vekillist[i]= StringVar()
            self.tclist[i] = StringVar()
            self.verginolist[i] = StringVar()
            self.telnolist[i] = StringVar()

            Entry(self.root, width=50, borderwidth=1,textvariable=self.islist[i]).grid(row=row,column=3, columnspan=20)
            Entry(self.root, width=50, borderwidth=1,textvariable=self.adreslist[i]).grid(row=row+1,column=3, columnspan=20)
            Entry(self.root, width=50, borderwidth=1,textvariable=self.vekillist[i]).grid(row=row+2,column=3, columnspan=20)
            Entry(self.root, width=50, borderwidth=1,textvariable=self.tclist[i]).grid(row=row+3,column=3, columnspan=20)
            Entry(self.root, width=50, borderwidth=1,textvariable=self.verginolist[i]).grid(row=row+4,column=3, columnspan=20)
            Entry(self.root, width=50, borderwidth=1,textvariable=self.telnolist[i]).grid(row=row+5,column=3, columnspan=20)
            

            try:
                self.islist[i].set(self.isimlistesi[i])
            except:
                self.islist[i].set("")
            try:
                self.adreslist[i].set(self.adreslistesi[i])
            except:
                self.adreslist[i].set("")
            try:
                self.vekillist[i].set(self.vekillistesi[i])
            except:
                self.vekillist[i].set("")
            try:
                self.tclist[i].set(self.tcnolistesi[i])
            except:
                self.tclist[i].set("")
            try:
                self.telnolist[i].set(self.telnolistesi[i])
            except:
                self.telnolist[i].set("")
            try:
                self.verginolist[i].set(self.verginolistesi[i])
            except:
                self.verginolist[i].set("")
            Label(self.root, text=str(i+1)+". Başvurulan Bilgileri :").grid(row=row-1, column=1,columnspan=2)
            Label(self.root, text="İsim Soyisim :").grid(row=row, column=1,columnspan=2)
            Label(self.root, text="Adresi :").grid(row=row+1, column=1,columnspan=2)
            Label(self.root, text="Vekili :").grid(row=row+2, column=1,columnspan=2)
            Label(self.root, text="TC No :").grid(row=row+3, column=1,columnspan=2)           
            Label(self.root, text="Vergi No :").grid(row=row+4, column=1,columnspan=2)     
            Label(self.root, text="Tel No :").grid(row=row+5, column=1,columnspan=2)     
            row=row+10
    def popupmsg(self,msg):
        popup = tk.Tk()
        popup.wm_title("!")
        label = ttk.Label(popup, text=msg, font=NORM_FONT)
        label.pack(side="top", fill="x", pady=20)
        B1 = ttk.Button(popup, text="Tamam", command = popup.destroy)
        B1.pack()
        popup.mainloop()
    def veriyaz(self):
        basvurantc=self.tc_entry.get()
        basvuranisim=self.basvuran_isim_entry.get()
        basvuranadres=self.basvuran_adres_entry.get()
        basvuranvekil=self.basvuran_vekil_entry.get()
        burono=self.buro_entry.get()
        basvuruno=self.basvuru_entry.get()
    
        
        

        htmldata=""
        for i in range(0,len(self.isimlistesi)):
            isim=self.islist[i].get()
            adres=self.adreslist[i].get()
            vekil=self.vekillist[i].get()
            tckimlik=self.tclist[i].get()
            vergino=self.verginolist[i].get()
            telno=self.telnolist[i].get()
            htmldata=htmldata+"""
            
            
            
                    İsim Soyisim"""+str(isim)+"""

                    Tc Kimlik Numarası"""+str(tckimlik)+"""
               
                    Adresi :"""+str(adres)+"""

                    Vekili : """+str(vekil)+"""

                    Vergi No : """+str(vergino)+"""

                    Telefon No : """+str(telno)+"""
            
            """
        


        html_str = """
       
        HUKUK UYUŞMAZLIKLARINDA DAVA ŞARTI ARABULUCULUK  İLK OTURUM TUTANAĞI
        

        Arabuluculuk Bürosu : """+str(self.burono)+"""

        Arabuluculuk Numarası :"""+str(basvuruno)+"""

        Başvuru Tarihi :"""+str()+"""



        ARABULUCUNUN 

        Adı ve Soyadı :
        T.C. Kimlik Numarası :
        Arabulucu Sicil Numarası :
        Adresi :

        BAŞVURUCU
        

        Adı ve Soyadı :"""+str(basvuranisim)+"""

        T.C. Kimlik Numarası :"""+str(basvurantc)+"""

    

        Adresi :"""+str(basvuranadres)+"""

        Vekili :"""+str(basvuranvekil)+"""

        
        KARŞI TARAF BİLGİLERİ
       
        """+htmldata+"""
"""
        
        doc= docx.Document()

        doc.add_paragraph(html_str)
        doc.save('Sonuc.docx')
        self.popupmsg("Başarıyla Word Haline Getirildi")

    def veritemizle(self):
        line_list = self.text1.get('1.0', 'end').split('\n')
            
        
        self.tc,self.isimsoyisim,self.adres,self.basvuranisimsoyisim,self.basvuranadres,self.hukukburo,self.basvurunumarasi,self.basvurutarihi,self.vekili,self.isimlistesi,self.adreslistesi,self.vekillistesi,self.verginolistesi,self.telnolistesi,self.tcnolistesi=Vericekici().vericek(line_list)
        
        self.tc_entry.set("")
        self.basvuran_isim_entry.set("")
        self.basvuran_adres_entry.set("")
        self.basvuran_vekil_entry.set("")
        self.buro_entry.set("")
        self.basvuru_entry.set("")
        self.basvurutarihi_entry.set("")


        self.islist=list(range(0,len(self.isimlistesi)))
        self.adreslist=list(range(0,len(self.isimlistesi)))
        self.vekillist=list(range(0,len(self.isimlistesi)))
        self.tclist=list(range(0,len(self.isimlistesi)))
        self.verginolist=list(range(0,len(self.isimlistesi)))
        self.telnolist=list(range(0,len(self.isimlistesi)))
        
        row=11
        for i in range(0,len(self.isimlistesi)):
            self.islist[i] = StringVar()
            self.adreslist[i]= StringVar()
            self.vekillist[i]= StringVar()
            self.tclist[i] = StringVar()
            self.verginolist[i] = StringVar()
            self.telnolist[i] = StringVar()

            Entry(self.root, width=50, borderwidth=1,textvariable=self.islist[i]).grid(row=row,column=3, columnspan=20)
            Entry(self.root, width=50, borderwidth=1,textvariable=self.adreslist[i]).grid(row=row+1,column=3, columnspan=20)
            Entry(self.root, width=50, borderwidth=1,textvariable=self.vekillist[i]).grid(row=row+2,column=3, columnspan=20)
            Entry(self.root, width=50, borderwidth=1,textvariable=self.tclist[i]).grid(row=row+3,column=3, columnspan=20)
            Entry(self.root, width=50, borderwidth=1,textvariable=self.verginolist[i]).grid(row=row+4,column=3, columnspan=20)
            Entry(self.root, width=50, borderwidth=1,textvariable=self.telnolist[i]).grid(row=row+5,column=3, columnspan=20)
            

            try:
                self.islist[i].set("")
            except:
                self.islist[i].set("")
            try:
                self.adreslist[i].set("")
            except:
                self.adreslist[i].set("")
            try:
                self.vekillist[i].set("")
            except:
                self.vekillist[i].set("")
            try:
                self.tclist[i].set("")
            except:
                self.tclist[i].set("")
            try:
                self.telnolist[i].set("")
            except:
                self.telnolist[i].set("")
            try:
                self.verginolist[i].set("")
            except:
                self.verginolist[i].set("")
            Label(self.root, text=str(i+1)+". Başvurulan Bilgileri :").grid(row=row-1, column=1,columnspan=2)
            Label(self.root, text="İsim Soyisim :").grid(row=row, column=1,columnspan=2)
            Label(self.root, text="Adresi :").grid(row=row+1, column=1,columnspan=2)
            Label(self.root, text="Vekili :").grid(row=row+2, column=1,columnspan=2)
            Label(self.root, text="TC No :").grid(row=row+3, column=1,columnspan=2)           
            Label(self.root, text="Vergi No :").grid(row=row+4, column=1,columnspan=2)     
            Label(self.root, text="Tel No :").grid(row=row+5, column=1,columnspan=2)     
            row=row+10        
    def ticaridavasontutanak(self):
        basvurantc=self.tc_entry.get()
        basvuranisim=self.basvuran_isim_entry.get()
        basvuranadres=self.basvuran_adres_entry.get()
        basvuranvekil=self.basvuran_vekil_entry.get()
        burono=self.buro_entry.get()
        basvuruno=self.basvuru_entry.get()
        
            
            

        htmldata=""
        for i in range(0,len(self.isimlistesi)):
            isim=self.islist[i].get()
            adres=self.adreslist[i].get()
            vekil=self.vekillist[i].get()
            tckimlik=self.tclist[i].get()
            vergino=self.verginolist[i].get()
            telno=self.telnolist[i].get()
            htmldata=htmldata+"""
                
                
                
                        İsim Soyisim"""+str(isim)+"""

                        Tc Kimlik Numarası"""+str(tckimlik)+"""
                
                        Adresi :"""+str(adres)+"""

                        Vekili : """+str(vekil)+"""

                        Vergi No : """+str(vergino)+"""

                        Telefon No : """+str(telno)+"""
                
                """
            


            html_str = """
        
        HUKUK UYUŞMAZLIKLARINDA DAVA ŞARTI ARABULUCULUK SON TUTANAĞI

        
        Arabuluculuk Bürosu				      : 	"""+self.hukukburo+"""
            Büro Dosya Numarası			      :	"""+basvuruno+"""
        Arabuluculuk Numarası			      :	************

        Arabulucunun					      :
                Adı ve Soyadı				                  :	**************
            T.C. Kimlik Numarası		                  :	**************
            Arabulucu Sicil Numarası                                   :     **************
                Adresi						      :	**************
                
        """+str(htmldata)+"""    

        Arabuluculuk Konusu Uyuşmazlık		      :	Ticari Uyuşmazlık
        Arabuluculuk Sürecinin Başladığı Tarih                    :	1.3.2019
        Arabuluculuk Sürecinin Bittiği Tarih   	                  :	21.4.2019

        Son Tutanağın Düzenlendiği Yer			      :	***************     
        Son Tutanağın Düzenlendiği Tarih		      :	21.4.2019

        Arabuluculuk Sonucu				      :

            Adı geçen taraflar ************************************** Toplantı Odasına geldiler.
        Taraflara arabuluculuğun temel ilkeleri, arabuluculuk süreci ve arabuluculuk süreci sonunda hazırlanan arabuluculuk son tutanağının hukuki ve mali yönlerden bütün sonuçları hakkında bilgi verildi.
        Taraflar söz alarak arabuluculuğun temel ilkelerini, arabuluculuk sürecini ve arabuluculuk süreci sonunda hazırlanan arabuluculuk son tutanağının hukuki ve mali yönlerden bütün sonuçlarını anladık dediler.
        Alternatif (1)
        Taraflar müzakereler sonucunda anlaşmaya varmışlardır. Taraflar anlaştıklarını beyan ettiler ve son tutanağa böyle geçsin dediler.
        Alternatif (2)
        Taraflar müzakereler sonucunda anlaşmaya varamamışlardır. Taraflar anlaşamadıklarını beyan ettiler ve son tutanağa böyle geçsin dediler. Tarafların üzerinde anlaşamadığı uyuşmazlık konusu/konuları …………………………dır.
        İşbu arabuluculuk son tutanağı iki sayfa ve üç nüsha olarak 6325 sayılı Hukuk Uyuşmazlıklarında Arabuluculuk Kanunu m. 17 ile 6102 sayılı Türk Ticaret Kanunu m. 5/A uyarınca hep birlikte imza altına alındı.



                                                İmzalar                                               
        Taraf 1		: 	*************			/
        Taraf 2		: 	************* Tic. Ltd. Şti.
                                        Adına Yetkili Temsilci **********	/
        Arabulucu		: 	*************			/


        """



        doc= docx.Document()

        doc.add_paragraph(html_str)
        doc.save('Sonuc.docx')
        self.popupmsg("Başarıyla Word Haline Getirildi")

    def ticaridavailktutanak(self):
        basvurantc=self.tc_entry.get()
        basvuranisim=self.basvuran_isim_entry.get()
        basvuranadres=self.basvuran_adres_entry.get()
        basvuranvekil=self.basvuran_vekil_entry.get()
        burono=self.buro_entry.get()
        basvuruno=self.basvuru_entry.get()
    
        
        

        htmldata=""
        for i in range(0,len(self.isimlistesi)):
            isim=self.islist[i].get()
            adres=self.adreslist[i].get()
            vekil=self.vekillist[i].get()
            tckimlik=self.tclist[i].get()
            vergino=self.verginolist[i].get()
            telno=self.telnolist[i].get()
            htmldata=htmldata+"""
            
            
            
                    İsim Soyisim"""+str(isim)+"""

                    Tc Kimlik Numarası"""+str(tckimlik)+"""
               
                    Adresi :"""+str(adres)+"""

                    Vekili : """+str(vekil)+"""

                    Vergi No : """+str(vergino)+"""

                    Telefon No : """+str(telno)+"""
            
            """
        


        html_str = """
        
    HUKUK UYUŞMAZLIKLARINDA DAVA ŞARTI ARABULUCULUK İLK OTURUM/AÇILIŞ TUTANAĞI

    
    Arabuluculuk Bürosu				      : 	"""+str(self.hukukburo)+"""
        Büro Dosya Numarası			      :	"""+str(basvuruno)+"""
    Arabuluculuk Numarası			      :	"""+str(basvuruno)+"""

    
    Arabulucunun					      :
            Adı ve Soyadı				                  :	**************
        T.C. Kimlik Numarası		                  :	**************
        Arabulucu Sicil Numarası                                   :     **************
            Adresi						      :	**************
            
    """+str(htmldata)+"""

    Arabuluculuk Konusu Uyuşmazlık		      :	Ticari Uyuşmazlık
    Arabuluculuk Sürecinin Başladığı Tarih                    :	1.3.2019

    İlk Oturum/Açılış Tutanağının Düzenlendiği Yer      :	***************     
    İlk Oturum/Açılış Tutanağının Düzenlendiği Tarih  :	11.3.2019

    Arabuluculuk Sonucu			  	      :
    
        Adı geçen taraflar ************************************** Toplantı Odasına geldiler.
    Alternatif (1)
    Taraflara arabuluculuğun temel ilkeleri, arabuluculuk süreci ve arabuluculuk süreci sonunda hazırlanan arabuluculuk son tutanağının ve anlaşma belgesinin hukuki ve mali yönlerden bütün sonuçları hakkında bilgi verildi.

    Alternatif (2)
    Taraflara arabuluculuğun temel ilkeleri olan, arabuluculuk sürecinin iradi olduğu; arabuluculuk sürecinde her iki tarafın da eşit haklara sahip olduğu; taraflarca aksi kararlaştırılmadıkça arabulucunun arabuluculuk faaliyeti çerçevesinde kendisine sunulan veya diğer bir şekilde elde ettiği bilgi ve belgeler ile diğer kayıtları gizli tutmakla yükümlü olduğu ve tarafların ve görüşmelere katılan diğer kişilerin de bu konudaki gizliliğe uymak zorunda olduğu; tarafların, arabulucunun veya arabuluculuğa katılanlar da dâhil üçüncü bir kişinin, uyuşmazlıkla ilgili hukuk davası açıldığında yahut tahkim yoluna başvurulduğunda, tarafların arabuluculuk sürecine katılma isteğini, arabuluculuk sürecinde taraflarca ileri sürülen görüşleri, önerileri ya da herhangi bir vakıanın veya iddianın kabulünü ve sadece arabuluculuk faaliyeti dolayısıyla hazırlanan belgeleri delil olarak ileri süremeyeceği ve bunlar hakkında tanıklık yapamayacağı hususları hakkında bilgi verildi.
    Taraflara arabuluculuk faaliyeti sonunda anlaşmaları hâlinde, arabuluculuk ücretinin, Arabuluculuk Asgari Ücret Tarifesinin eki Arabuluculuk Ücret Tarifesinin İkinci Kısmına göre aksi kararlaştırılmadıkça taraflarca eşit şekilde karşılanacağı, bu durumda ücretin Tarifenin Birinci Kısmında belirlenen iki saatlik ücret tutarından az olamayacağı; arabuluculuk faaliyeti sonunda iki saatten az süren görüşmeler sonunda tarafların anlaşamamaları hâllerinde iki saatlik ücret tutarının Tarifenin Birinci Kısmına göre Adalet Bakanlığı bütçesinden ödeneceği, iki saatten fazla süren görüşmeler sonunda tarafların anlaşamamaları hâlinde ise iki saati aşan kısma ilişkin ücretin aksi kararlaştırılmadıkça taraflarca eşit şekilde uyuşmazlığın konusu dikkate alınarak Tarifenin Birinci Kısmına göre karşılanacağı, Adalet Bakanlığı bütçesinden ödenen ve taraflarca karşılanan arabuluculuk ücretinin yargılama giderlerinden sayılacağı hususları hakkında bilgi verildi.
    Taraflara arabulucunun görevini özenle, tarafsız bir biçimde ve şahsen yerine getireceği, arabulucunun taraflar arasında eşitliği gözetmekle yükümlü olduğu, arabuluculuk müzakerelerine tarafların bizzat, kanuni temsilcileri veya vekâletnamesinde özel yetki bulunan avukatları aracılığıyla katılabileceği, arabuluculuk sürecinde arabulucunun rolünün, hâkim veya hakem olmadığı, kimin haklı ya da haksız olduğu konusunda karar vermeyeceği, yargısal bir yetkinin kullanımı olarak sadece hâkim tarafından yapılabilecek işlemleri yapamayacağı, taraflara hukuki tavsiyelerde bulunamayacağı, tarafların çözüm üretemediklerinin ortaya çıkması hâlinde arabulucunun bir çözüm önerisinde bulunabileceği, bununla birlikte bir çözüm önerisi ya da öneriler kataloğu geliştirip bunu taraflara empoze edemeyeceği, müzakereler sırasında geliştirilen bir çözüm önerisi üzerinde anlaşmaya varmaları için tarafları zorlayamayacağı; bununla birlikte, yaşanılan uyuşmazlık ile ilgili çözüm seçeneklerini üreterek bir anlaşmaya ulaşabilmelerinde taraflara yardımcı olacak iletişimin ortamını sağlayacağı, bilgileri dâhilinde ve onay vermeleri hâlinde taraflarla ayrı ayrı veya birlikte görüşebileceği ve iletişim kurabileceği, arabulucu olarak tarafsız bir konumda olduğu, arabuluculuk sürecinin sonunda her iki tarafın da kabul edeceği bir anlaşmaya varılamaması hâlinde açılabilecek olası bir davada, daha sonra avukat olarak görev üstlenemeyeceği, arabuluculuk bürosuna başvurulmasından son tutanağın düzenlendiği tarihe kadar geçen sürede zamanaşımının duracağı ve hak düşürücü sürenin işlemeyeceği, arabuluculuk sürecinin sonunda her iki tarafın da kabul edeceği bir anlaşmaya varılamaması hâlinde yargı organlarına başvuru haklarının bulunduğu hususları hakkında bilgi verildi.
    Taraflara arabuluculuk sürecinde düzenlenecek oturum tutanaklarına ve sürecin sonunda düzenlenecek son tutanağa, oturumların ve faaliyetin sonuçlanması dışında hangi hususların yazılacağına tarafların karar vereceği, aynı şekilde arabuluculuk sürecinin sonunda varılan anlaşmanın kapsamının taraflarca belirleneceği, anlaşma belgesi düzenlenmesi hâlinde bu belgenin taraflar veya avukatları ve arabulucu tarafından imzalanacağı, tarafların bu anlaşma belgesinin icra edilebilirliğine ilişkin mahkemeden şerh verilmesini talep edebileceği ve bu şerhi içeren anlaşmanın ilâm niteliğinde belge sayılacağı, taraflar ve avukatları ile arabulucunun birlikte imzaladıkları anlaşma belgesinin icra edilebilirlik şerhi aranmaksızın ilâm niteliğinde belge sayılacağı, arabuluculuk faaliyeti sonunda anlaşmaya varılması hâlinde üzerinde anlaşılan hususlar hakkında taraflarca dava açılamayacağı hususları hakkında bilgi verildi.
    Ayrıca, taraflara, kendilerinden arabuluculuk sürecinde birbirlerine karşı “siz”li hitap şeklini kullanmalarının ve söz verildiği zaman, sırayla ve sözleri kesilmeden konuşmalarının beklendiği, birbirlerinin sözünü kesmelerinin, söz veya hareketle diğer tarafı övmelerinin veya tahkir etmelerinin yasak olduğu, daha sonra eklemek istedikleri hususlar hakkında kendilerine konuşma olanağı tanınacağı, arabulucu tarafından da kendilerine sorular sorulabileceği hususları belirtilmiş; arabuluculuk sürecinde olabildiğince açık ve dürüst olunmasının ve işbirliği hâlinde hareket edilmesinin önemi vurgulanmış; arabuluculuk sürecinde belirtilen kurullara uymayı kabul edip etmedikleri kendilerine sorulmuştur.
    Taraflar söz alarak arabuluculuğun temel ilkelerini, arabuluculuk sürecini ve arabuluculuk süreci sonunda hazırlanan arabuluculuk son tutanağının ve anlaşma belgesinin hukuki ve mali yönlerden bütün sonuçlarını anladık ve arabuluculuk sürecinde belirtilen kurullara uymayı kabul ediyoruz dediler.
    Taraflar müzakerelere başlamışlardır.
    İşbu arabuluculuk ilk oturum/açılış tutanağı dört sayfa ve üç nüsha olarak 6325 sayılı Hukuk Uyuşmazlıklarında Arabuluculuk Kanunu m. 11, m. 15 ve m. 16 ile 6102 sayılı Türk Ticaret Kanunu m. 5/A uyarınca hep birlikte imza altına alındı.


                                                İmzalar                                               
    Taraf 1		: 	*************			/
    Taraf 2		: 	************* Tic. Ltd. Şti.
                                    Adına Yetkili Temsilci **********	/
    Arabulucu		: 	*************			/


    """
        
        doc= docx.Document()

        doc.add_paragraph(html_str)
        doc.save('Sonuc.docx')
        self.popupmsg("Başarıyla Word Haline Getirildi")
  


    def ishukukubilgilendirmetutanagi(self):
        basvurantc=self.tc_entry.get()
        basvuranisim=self.basvuran_isim_entry.get()
        basvuranadres=self.basvuran_adres_entry.get()
        basvuranvekil=self.basvuran_vekil_entry.get()
        burono=self.buro_entry.get()
        basvuruno=self.basvuru_entry.get()
    
        
        

        htmldata=""
        for i in range(0,len(self.isimlistesi)):
            isim=self.islist[i].get()
            adres=self.adreslist[i].get()
            vekil=self.vekillist[i].get()
            tckimlik=self.tclist[i].get()
            vergino=self.verginolist[i].get()
            telno=self.telnolist[i].get()
            htmldata=htmldata+"""
            
            
            
                    İsim Soyisim"""+str(isim)+"""

                    Tc Kimlik Numarası"""+str(tckimlik)+"""
               
                    Adresi :"""+str(adres)+"""

                    Vekili : """+str(vekil)+"""

                    Vergi No : """+str(vergino)+"""

                    Telefon No : """+str(telno)+"""
            
            """
        


        html_str = """
        
            -ArabuluculukSürecineİlişkinBilgilendirmeTutanağı-
        ..... /..... /..... tarihindeaşağıda ad veunvanlarıyazılıtaraflarbirlikte/ayrıayrı 6325 sayılıHukukUyuşmazlıklarındaArabuluculukKanunugereğincearabuluculuksürecineilişkinolarakaşağıdakikonularda,davetmektubuekindeyazılıolarakvetoplantıdakatılımlarıylasözlüolarakbilgilendirilmişlerdir:
        GENEL BİLGİLER:
        	Taraflar, arabulucuyabaşvurmak, sürecidevamettirmek, sonuçlandırmakveyabusüreçtenvazgeçmekkonusundatamamenserbesttirler. (Ek cümle:6/12/20187155/22 md.)  Şukadarkidavaşartıolarakarabuluculuğailişkin 18/A maddesihükmüsaklıdır. Taraflarzorlabusüreciniçinedâhiledilemeyeceklerigibi her aşamadauyuşmazlığıarabuluculukyoluylasonuçlandırmaktan davazgeçebilirler. Taraflar, gerekarabulucuyabaşvururkengereksesüreçboyuncaeşithaklarasahiptirler. Taraflardanbiriarabuluculuksürecinindışındabırakılamayacağıgibisözhakkıdadiğerinegörekısıtlanamaz.Arabuluculukyoluylauyuşmazlığınçözümügönüllülükesasınadayalıdır.
        	Arabulucutaraflararasındakihukukiuyuşmazlığınçözümündetarafsızvebağımsızbirüçüncükişiolarakyeralırvetaraflararasındakiiletişimortamınıkolaylaştırarakkendiçözümlerinikendilerininüretmelerikonusundaonlarayardımcıolur. Taraflarınçözümüretemediklerininortayaçıkmasıhâlindearabulucubirçözümönerisinde de bulunabilir.
        	Taraflarcaaksikararlaştırılmadıkçaarabuluculukgörüşmelerindegizlilikilkesineuyulmasıesastır. Taraflarcaaksikararlaştırılmadıkçaarabulucu, taraflarvearabuluculuğakatılanüçüncükişiler, taraflarcayapılanarabuluculukdavetiveyabirtarafınarabuluculukfaaliyetinekatılmaisteği; uyuşmazlığınarabuluculukyoluilesonaerdirilmesiiçintaraflarcailerisürülengörüşveteklifler; arabuluculukfaaliyetiesnasındataraflarcailerisürülenönerilerveyaherhangibirvakıaveyaiddianınkabulü; sadecearabuluculukfaaliyetidolayısıylahazırlananbelgeleridelilolarakilerisüremezvebunlarhakkındatanıklıkyapamaz.Buyükümlülüğeaykırıhareketederekbirkişininhukukenkorunanmenfaatininzarargörmesinenedenolankişialtıayakadarhapiscezasıilecezalandırılır.	
        	Arabuluculuksüreci,davaaçılmadanöncearabulucuyabaşvuruhalinde, taraflarınilktoplantıyadavetedilmelerivetaraflarlaarabulucuarasındasürecindevamettirilmesikonusundaanlaşmayavarılıpbudurumunbirtutanaklabelgelendirildiğitarihtenitibarenişlemeyebaşlar. Davaaçılmasındansonraarabulucuyabaşvuruhalindeisebusüreç,mahkemenintaraflarıarabuluculuğadavetinintaraflarcakabuledilmesiveyataraflarınarabulucuyabaşvurmakonusundaanlaşmayavardıklarınımahkemeyeyazılıolarakbeyanettikleriyadaduruşmadabubeyanlarınıntutanağageçirildiğitarihtenitibarenişlemeyebaşlar. (Davaşartıarabuluculuksüreci, adliyearabuluculukbürosunabaşvurulanhallerde, başvurununyapıldığıtarihtebaşlar.)
        	Arabuluculuksürecininbaşlamasındansonaermesinekadargeçensüre, zamanaşımıvehakdüşürücüsürelerinhesaplanmasındadikkatealınmaz.
        	Arabulucu, busıfatlagörevyaptığıuyuşmazlıkileilgiliolarakaçılan da- vada, dahasonraavukatolarakgörevüstlenemez.
        	Arabuluculukyoluylaçözümlenenhukukiuyuşmazlıklarkonusundaarabuluculukfaaliyetisonucundaanlaşmayavarılmasıhalindeanlaşmabelgesidüzenlenir. Bu anlaşmanıngereklerinintaraflardanherhangibiritarafındanyerinegetirilmemesidurumundadiğertarafarabuluculukanlaşmabelgesiniyetkilimahkemeyeibrazederekicraedilebilirlikşerhiverilmesinitalepedebilir, bubelgeyetkilimahkemetarafındanşerhverilmesiilebirlikteilâmniteliğindebelgevasfınıkazanarakmahkemekararıgibiicraedilir.Arabulucuileavukatlarvetaraflarcabirlikteimzaaltınaalınanarabuluculukanlaşmalarıdoğrudanilamniteliğindebelgevasfınıhaizolur. 
        	DAVA ŞARTI ARABULUCULUK HAKKINDA BİLGİLER:
        	İŞ HUKUKUNDA DAVA ŞARTI OLARAK DÜZENLENEN ARABULUCULUK SÜRECİ 
        	İş Mahkemeleri Kanunu m. 5/1,a hükmü uyarınca, 5953 sayılı Basın Mesleğinde Çalışanlarla Çalıştıranlar Arasındaki Münasebetlerin Tanzimi Hakkında Kanuna tâbi gazeteciler, 854 sayılı Deniz İş Kanununa tâbi gemiadamları, 22/5/2003 tarihli ve 4857 sayılı İş Kanununa veya 11/1/2011 tarihli ve 6098 sayılı Türk Borçlar Kanununun İkinci Kısmının Altıncı Bölümünde düzenlenen hizmet sözleşmelerine tâbi işçiler ile işveren veya işveren vekilleri arasında, iş ilişkisi nedeniyle sözleşmeden veya kanundan doğan her türlü hukuk uyuşmazlıklarına iş mahkemelerinde bakılacaktır.
        	İşçiveişverenarasındakiuyuşmazlıklarda,davaşartıolarakarabuluculuk, 7036 SayılıİşMahkemeleriKanununun3.maddesinde düzenlemealtınaalınmıştır.Bunagöre;
        	(1) Kanuna, bireysel veya toplu iş sözleşmesine dayanan işçi veya işveren alacağı ve tazminatı ile işe iade talebiyle açılan davalarda, arabulucuya başvurulmuş olması dava şartıdır. 
        	(2) Davacı, arabuluculuk faaliyeti sonunda anlaşmaya varılamadığına ilişkin son tutanağın aslını veya arabulucu tarafından onaylanmış bir örneğini dava dilekçesine eklemek zorundadır. Bu zorunluluğa uyulmaması hâlinde mahkemece davacıya, son tutanağın bir haftalık kesin süre içinde mahkemeye sunulması gerektiği, aksi takdirde davanın usulden reddedileceği ihtarını içeren davetiye gönderilir. İhtarın gereği yerine getirilmez ise dava dilekçesi karşı tarafa tebliğe çıkarılmaksızın davanın usulden reddine karar verilir. Arabulucuya başvurulmadan dava açıldığının anlaşılması hâlinde herhangi bir işlem yapılmaksızın davanın, dava şartı yokluğu sebebiyle usulden reddine karar verilir. 
        	(3) İş kazası veya meslek hastalığından kaynaklanan maddi ve manevi tazminat ile bunlarla ilgili tespit, itiraz ve rücu davaları hakkında birinci fıkra hükmü uygulanmaz.  
        	(4) Arabuluculuk Daire Başkanlığı, sicile kayıtlı arabuluculardan bu madde uyarınca arabuluculuk yapmak isteyenleri, varsa uzmanlık alanlarını da belirterek, görev yapmak istedikleri adli yargı ilk derece mahkemesi adalet komisyonlarına göre listeler ve listeleri ilgili komisyon başkanlıklarına bildirir. Komisyon başkanlıkları, bu listeleri kendi yargı çevrelerindeki arabuluculuk bürolarına, arabuluculuk bürosu kurulmayan yerlerde ise görevlendirecekleri sulh hukuk mahkemesi yazı işleri müdürlüğüne gönderir. (5) Başvuru karşı tarafın, karşı taraf birden fazla ise bunlardan birinin yerleşim yerindeki veya işin yapıldığı yerdeki arabuluculuk bürosuna, arabuluculuk bürosu kurulmayan yerlerde ise görevlendirilen yazı işleri müdürlüğüne yapılır. 
        	(6) Arabulucu, komisyon başkanlıklarına bildirilen listeden büro tarafından belirlenir. Ancak tarafların listede yer alan herhangi bir arabulucu üzerinde anlaşmaları hâlinde bu arabulucu görevlendirilir. 
        	(7) Başvuran taraf, kendisine ve elinde bulunması hâlinde karşı tarafa ait her türlü iletişim bilgisini arabuluculuk bürosuna verir. Büro, tarafların resmi kayıtlarda yer alan iletişim bilgilerini araştırmaya da yetkilidir. İlgili kurum ve kuruluşlar, büro tarafından talep edilen bilgi ve belgeleri vermekle yükümlüdür. 
        	(8) Taraflara ait iletişim bilgileri, görevlendirilen arabulucuya büro tarafından verilir. Arabulucu bu iletişim bilgilerini esas alır, ihtiyaç duyduğunda kendiliğinden araştırma da yapabilir. Elindeki bilgiler itibarıyla her türlü iletişim vasıtasını kullanarak görevlendirme konusunda tarafları bilgilendirir ve ilk toplantıya davet eder. Bilgilendirme ve davete ilişkin işlemlerini belgeye bağlar.
        	 (9) Arabulucu, görevlendirmeyi yapan büronun yetkili olup olmadığını kendiliğinden dikkate alamaz. Karşı taraf en geç ilk toplantıda, yerleşim yeri ve işin yapıldığı yere ilişkin belgelerini sunmak suretiyle arabuluculuk bürosunun yetkisine itiraz edebilir. Bu durumda arabulucu, dosyayı derhâl ilgili sulh hukuk mahkemesine gönderilmek üzere büroya teslim eder. Mahkeme, harç alınmaksızın dosya üzerinden yapacağı inceleme sonunda yetkili büroyu kesin olarak karara bağlar ve dosyayı büroya iade eder. Mahkeme kararı büro tarafından 11/2/1959 tarihli ve 7201 sayılı Tebligat Kanunu hükümleri uyarınca taraflara tebliğ edilir. 
            Yetki itirazının reddi durumunda aynı arabulucu yeniden görevlendirilir ve onuncu fıkrada belirtilen süreler yeni görevlendirme tarihinden başlar. Yetki itirazının kabulü durumunda ise kararın tebliğinden itibaren bir hafta içinde yetkili büroya başvurulabilir. Bu takdirde yetkisiz büroya başvurma tarihi yetkili büroya başvurma tarihi olarak kabul edilir. Yetkili büro, altıncı fıkra uyarınca arabulucu görevlendirir. 
        	(10) Arabulucu, yapılan başvuruyu görevlendirildiği tarihten itibaren üç hafta içinde sonuçlandırır. Bu süre zorunlu hâllerde arabulucu tarafından en fazla bir hafta uzatılabilir. 
        	(11) Arabulucu, taraflara ulaşılamaması, taraflar katılmadığı için görüşme yapılamaması yahut yapılan görüşmeler sonucunda anlaşmaya varılması veya varılamaması hâllerinde arabuluculuk faaliyetini sona erdirir ve son tutanağı düzenleyerek durumu derhâl arabuluculuk bürosuna bildirir. 
        	12) Taraflardan birinin geçerli bir mazeret göstermeksizin ilk toplantıya katılmaması sebebiyle arabuluculuk faaliyetinin sona ermesi durumunda toplantıya katılmayan taraf, son tutanakta belirtilir ve bu taraf davada kısmen veya tamamen haklı çıksa bile yargılama giderinin tamamından sorumlu tutulur. Ayrıca bu taraf lehine vekâlet ücretine hükmedilmez. Her iki tarafın da ilk toplantıya katılmaması sebebiyle sona eren arabuluculuk faaliyeti üzerine açılacak davalarda tarafların yaptıkları yargılama giderleri kendi üzerlerinde bırakılır. 
        	(13) Tarafların arabuluculuk faaliyeti sonunda anlaşmaları hâlinde, arabuluculuk ücreti, Arabuluculuk Asgari Ücret Tarifesinin eki Arabuluculuk Ücret Tarifesinin İkinci Kısmına göre aksi kararlaştırılmadıkça taraflarca eşit şekilde karşılanır. Bu durumda ücret, Tarifenin Birinci Kısmında belirlenen iki saatlik ücret tutarından az olamaz. İşe iade talebiyle yapılan görüşmelerde tarafların anlaşmaları durumunda, arabulucuya ödenecek ücretin belirlenmesinde işçiye işe başlatılmaması hâlinde ödenecek tazminat miktarı ile çalıştırılmadığı süre için ödenecek ücret ve diğer haklarının toplamı, Tarifenin İkinci Kısmı uyarınca üzerinde anlaşılan miktar olarak kabul edilir. 
        	(14) Arabuluculuk faaliyeti sonunda taraflara ulaşılamaması, taraflar katılmadığı için görüşme yapılamaması veya iki saatten az süren görüşmeler sonunda tarafların anlaşamamaları hâllerinde, iki saatlik ücret tutarı Tarifenin Birinci Kısmına göre Adalet Bakanlığı bütçesinden ödenir. İki saatten fazla süren görüşmeler sonunda tarafların anlaşamamaları hâlinde ise iki saati aşan kısma ilişkin ücret aksi kararlaştırılmadıkça taraflarca eşit şekilde Tarifenin Birinci Kısmına göre karşılanır. Adalet Bakanlığı bütçesinden ödenen ve taraflarca karşılanan arabuluculuk ücreti, yargılama giderlerinden sayılır. 
        	(15) Asıl işveren-alt işveren ilişkisinin varlığı hâlinde işe iade talebiyle arabulucuya başvurulduğunda, anlaşmanın gerçekleşebilmesi için işverenlerin arabuluculuk görüşmelerine birlikte katılmaları ve iradelerinin birbirine uygun olması aranır. 
        	Ayrıca;“Arabuluculuk faaliyeti sonunda tarafların, işçinin işe başlatılması konusunda anlaşmaları hâlinde;  a) İşe başlatma tarihini, b) Üçüncü fıkrada düzenlenen ücret ve diğer hakların parasal miktarını, c) İşçinin işe başlatılmaması durumunda ikinci fıkrada düzenlenen tazminatın parasal miktarını, belirlemeleri zorunludur. Aksi takdirde anlaşma sağlanamamış sayılır ve son tutanak buna göre düzenlenir. İşçinin kararlaştırılan tarihte işe başlamaması hâlinde fesih geçerli hâle gelir ve işveren sadece bunun hukuki sonuçları ile sorumlu olur”
        	(16) Bu madde uyarınca arabuluculuk bürosu tarafından yapılması gereken zaruri giderler; arabuluculuk faaliyeti sonunda anlaşmaya varılması hâlinde anlaşma uyarınca taraflarca ödenmek, anlaşmaya varılamaması hâlinde ise ileride haksız çıkacak taraftan tahsil olunmak üzere Adalet Bakanlığı bütçesinden karşılanır. 
        	(17) Arabuluculuk bürosuna başvurulmasından son tutanağın düzenlendiği tarihe kadar geçen sürede zamanaşımı durur ve hak düşürücü süre işlemez. 
        	(18) Arabuluculuk görüşmelerine taraflar bizzat, kanuni temsilcileri veya avukatları aracılığıyla katılabilirler. İşverenin yazılı belgeyle yetkilendirdiği çalışanı da görüşmelerde işvereni temsil edebilir ve son tutanağı imzalayabilir. 
        	(19) Arabuluculuk görüşmeleri, taraflarca aksi kararlaştırılmadıkça, arabulucuyu görevlendiren büronun bağlı bulunduğu adli yargı ilk derece mahkemesi adalet komisyonunun yetki alanı içinde yürütülür. 
        	(20) 13/6/1952 tarihli ve 5953 sayılı Basın Mesleğinde Çalışanlarla Çalıştıranlar Arasındaki Münasebetlerin Tanzimi Hakkında Kanunda düzenlenen gazeteci ile 20/4/1967 tarihli ve 854 sayılı Deniz İş Kanununda düzenlenen gemiadamı, bu madde kapsamında işçi sayılır. 
        	(21) Bu maddede hüküm bulunmayan hâllerde niteliğine uygun düştüğü ölçüde 7/6/2012 tarihli ve 6325 sayılı Hukuk Uyuşmazlıklarında Arabuluculuk Kanunu hükümleri uygulanır. 
        	(22) Arabuluculuğa başvuru usulü, arabulucunun görevlendirilmesi ve arabuluculuk görüşmelerine ilişkin diğer hususlar Adalet Bakanlığınca yürürlüğe konulan yönetmelikle belirlenir.
        	Taraflararabuluculuksüreçveesaslarıileanlaşmanınhukukiniteliğiileilgili, işbuimzayakonubelgeiçindekiaçıklamalarıanladıklarınıvearabulucununtarafsızlığındanşüpheedilmesinigerektirecekhiçbirhalveşartınolmadığındanbahislearabuluculuksürecinebaşlamakistedikleriniifadeettiler. 
        	İşbututanaküçnüshaolarakdüzenlenmiştir. ..... /..... /.....
        (Ad,soyad, TC Kimlik no.veimza)	(Ad, soyad, TC Kimlik no. veimza)
            Taraf	                                                                    Taraf
            """+str(basvuranisim)+str(basvurantc)+"""
            ArabulucuAv.Arb.DilekYumrutaş
        ArabuluculukSicil No: 83 ...........

        Bu Tutanaktarafsayısındanbirfazlanüshaolarakdüzenlenir; taraflardavearabulucununkendisindekalır.

        """
        
        doc= docx.Document()

        doc.add_paragraph(html_str)
        doc.save('Sonuc.docx')
        self.popupmsg("Başarıyla Word Haline Getirildi")


    def arabuluculukbelirlemetutanagi(self):
        basvurantc=self.tc_entry.get()
        basvuranisim=self.basvuran_isim_entry.get()
        basvuranadres=self.basvuran_adres_entry.get()
        basvuranvekil=self.basvuran_vekil_entry.get()
        burono=self.buro_entry.get()
        basvuruno=self.basvuru_entry.get()
    
        
        

        htmldata=""
        for i in range(0,len(self.isimlistesi)):
            isim=self.islist[i].get()
            adres=self.adreslist[i].get()
            vekil=self.vekillist[i].get()
            tckimlik=self.tclist[i].get()
            vergino=self.verginolist[i].get()
            telno=self.telnolist[i].get()
            htmldata=htmldata+"""
                
                
                
                        İsim Soyisim"""+str(isim)+"""

                        Tc Kimlik Numarası"""+str(tckimlik)+"""
                
                        Adresi :"""+str(adres)+"""

                        Vekili : """+str(vekil)+"""

                        Vergi No : """+str(vergino)+"""

                        Telefon No : """+str(telno)+"""
                
                """
            


            html_str = """
        
                            ARABULUCU BELİRLEME TUTANAĞI

    BAŞVURAN	: """+str(basvuranisim)+"""
    ADRESİ		: """+str(basvuranadres)+"""
    İLETİŞİM	TEL:"""+str(telno)+"""	TEL(CEP):"""+str(telno)+""".	FAKS:……………..	
            E-POSTA:………………

    MUHATAP	:
    """+str(htmldata)+"""
    KONU		: İşçilik Alacakları

        Hukuk Uyuşmazlıklarında Arabuluculuk Kanunu ve İş Mahkemeleri Kanunu gereğince dava şartı arabuluculuk hükümleri çerçevesinde uyuşmazlığımızın Arabuluculuk yoluyla çözümü için tarafımızdan Arabuluculu olarak  ….. sicil numaralı …………………………. belirlenmiştir. 

        Arabuluculuk Bürosu’ndaki başvuru işlemlerinin bu şekilde yapılmasını talep ederiz. 12.12.2018


    Başvuran  				                       Muhatap 				            Arabulucu
    """+str(basvuranisim)+"""	          

    """
        
        doc= docx.Document()

        doc.add_paragraph(html_str)
        doc.save('Sonuc.docx')
        self.popupmsg("Başarıyla Word Haline Getirildi")


    def ticarialabuluculukdavetmektubu(self):
        basvurantc=self.tc_entry.get()
        basvuranisim=self.basvuran_isim_entry.get()
        basvuranadres=self.basvuran_adres_entry.get()
        basvuranvekil=self.basvuran_vekil_entry.get()
        burono=self.buro_entry.get()
        basvuruno=self.basvuru_entry.get()
    
        
        

        htmldata=""
        for i in range(0,len(self.isimlistesi)):
            isim=self.islist[i].get()
            adres=self.adreslist[i].get()
            vekil=self.vekillist[i].get()
            tckimlik=self.tclist[i].get()
            vergino=self.verginolist[i].get()
            telno=self.telnolist[i].get()
            htmldata=htmldata+"""
                
                
                
                        İsim Soyisim"""+str(isim)+"""

                        Tc Kimlik Numarası"""+str(tckimlik)+"""
                
                        Adresi :"""+str(adres)+"""

                        Vekili : """+str(vekil)+"""

                        Vergi No : """+str(vergino)+"""

                        Telefon No : """+str(telno)+"""
                
                """
            


            html_str = """
        
                        TİCARİ UYUŞMAZLIKLARDA DAVA ŞARTI ARABULUCULUK SÜRECİNE DAVET MEKTUBU

    Sayın """+str(basvuranisim)+""" ,
    ……………………….tarafından ………………….. Arabuluculuk Bürosuna yapılan başvuru üzerine UYAP Arabulucu Portal tarafından görevlendirilmiş T.C. Adalet Bakanlığı’ndaki resmi sicile kayıtlı ……..sicil numaralı arabulucuyum.

    ........................ konudaki hukuki uyuşmazlığınızı barışçıl olarak arabuluculuk yoluyla çözebilmeniz için bu davet yazısını yazıyoruz.

    Hukuki uyuşmazlığınızın 6325 Sayılı Hukuk Uyuşmazlıklarında Arabuluculuk Kanunu kapsamında tarafların üzerinde serbestçe tasarruf edebileceği iş ve işlemlerden doğan özel hukuk uyuşmazlığı olduğu anlaşılmaktadır
    6102 sayılı Türk Ticaret Kanunu uyarınca ticari davalardan, konusubirmiktarparanınödenmesiolanalacakvetazminattaleplerihakkındadavaaçılmadanöncearabulucuyabaşvurulmuşolmasıdavaşartıdır (TTK. m. 5A/1).
    Arabuluculuk sürecinin daha verimli geçmesi için, arabuluculukla ilgili şu hususları bilgilerinize sunmak isterim:
    a)Arabulucu taraflar arasındaki hukuki uyuşmazlığın çözümünde tarafsız ve bağımsız                                                                    bir üçüncü kişi olarak yer alır ve  taraflar arasındaki iletişim ortamını kolaylaştırarak kendi çözümlerini kendilerinin üretmeleri konusunda onlara yardımcı olur. Tarafların çözüm üretemediklerinin ortaya çıkması halinde arabulucu bir çözüm önerisinde de bulunabilir.
    b)Arabuluculuk yoluyla uyuşmazlığın çözümü gönüllülük esasına dayalıdır. Taraflar, süreci devam ettirmek, sonuçlandırmak veya bu süreçten vazgeçmek konusunda serbesttirler.
    c)Arabuluculuk yoluyla uyuşmazlığın çözümü ekonomik, sosyal ve psikolojik bakımdan faydalıdır. Arabuluculuksürecitaraflararasındakiilişkilerinkorunmasınayardımcıolurvetoplumsalbarışahizmeteder.
    ç)	Taraflarca aksi kararlaştırılmadıkça arabuluculuk görüşmelerinde gizlilik ilkesine uyulması esastır. Bu durum, ticari sırlarınızın korunmasını sağlayacağı gibi, ticari itibarınızın zarar görmesini de engelleyecektir.
    Arabuluculuk bürosuna başvurulmasından son tutanağın düzenlendiği tarihe kadar geçen sürede zamanaşımı durur ve hak düşürücü süre işlemez (HUAK m. 18A/15).
    Dava açılmadan önce ihtiyati tedbir kararı verilmesi hâlinde 6100 sayılı Kanunun 397 nci maddesinin birinci fıkrasında, ihtiyati haciz kararı verilmesi hâlinde ise 9/6/1932 tarihli ve 2004 sayılı İcra ve İflas Kanununun 264 üncü maddesinin birinci fıkrasında düzenlenen dava açma süresi, arabuluculuk bürosuna başvurulmasından son tutanağın düzenlendiği tarihe kadar işlemez (HUAK m. 18A/16).
    Davacı, arabuluculuk faaliyeti sonunda anlaşmaya varılamadığına ilişkin son tutanağın aslını veya arabulucu tarafından onaylanmış bir örneğini dava dilekçesine eklemek zorundadır. Bu zorunluluğa uyulmaması hâlinde mahkemece davacıya, son tutanağın bir haftalık kesin süre içinde mahkemeye sunulması gerektiği, aksi takdirde davanın usulden reddedileceği ihtarını içeren davetiye gönderilir. İhtarın gereği yerine getirilmez ise dava dilekçesi karşı tarafa tebliğe çıkarılmaksızın davanın usulden reddine karar verilir. Arabulucuya başvurulmadan dava açıldığının anlaşılması hâlinde herhangi bir işlem yapılmaksızın davanın, dava şartı yokluğu sebebiyle usulden reddine karar verilir (HUAK m. 3/2).
    Arabulucu, yapılan başvuruyu görevlendirildiği tarihten itibaren altı hafta içinde sonuçlandırır. Bu süre                                 zorunlu hâllerde arabulucu tarafından en fazla iki hafta uzatılabilir (TTK m. 5A/2).
    Arabulucu, taraflara ulaşılamaması, taraflar katılmadığı için görüşme yapılamaması yahut yapılan görüşmeler sonucunda anlaşmaya varılması veya varılamaması hâllerinde arabuluculuk	faaliyetini	sona	erdirir	ve	son	tutanağı	düzenleyerek durumu derhâl arabuluculuk bürosuna bildirir (HUAK m.18A/10).
    Tarafların arabuluculuk faaliyeti sonunda anlaşmaları hâlinde, arabuluculuk ücreti, Arabuluculuk Asgari Ücret Tarifesinin eki Arabuluculuk Ücret Tarifesinin İkinci Kısmına göre aksi kararlaştırılmadıkça taraflarca eşit şekilde karşılanır. Bu durumda ücret, Tarifenin Birinci Kısmında belirlenen iki saatlik ücret tutarından az olamaz (HUAK m. 18A/12).
    Arabuluculuk faaliyeti sonunda taraflara ulaşılamaması, taraflar katılmadığı için görüşme yapılamaması veya iki saatten az süren görüşmeler sonunda tarafların anlaşamamaları hâllerinde, iki saatlik ücret tutarı Tarifenin Birinci Kısmına göre Adalet Bakanlığı bütçesinden ödenir. İki saatten fazla süren görüşmeler sonunda tarafların anlaşamamaları hâlinde ise iki saati aş         an kısma ilişkin ücret aksi kararlaştırılmadıkça taraflarca eşit şekilde Tarifenin Birinci Kısmına göre karşılanır. Adalet Bakanlığı bütçesinden ödenen ve taraflarca karşılanan arabuluculuk ücreti, yargılama giderlerinden sayılır (HUAK m.18A/13).
    Arabuluculuk müzakerelerine taraflar bizzat, kanuni temsilcileri veya avukatları aracılığıyla katılabilirler. Uyuşmazlığın çözümüne katkı sağlayabilecek uzman kişiler de müzakerelerde hazır bulundurulabilir (HUAK m. 15/6).
    Arabuluculuk görüşmeleri, taraflarca aksi kararlaştırılmadıkça, arabulucuyu görevlendiren büronun bağlı bulunduğu adli yargı ilk derece mahkemesi adalet komisyonunun yetki alanı içinde yürütülür (HUAK m. 18A/17).

    Uyuşmazlığın tarafları olarak sizlerle yapacağımız ilk toplantı …………. Tarihinde saat…..’de …………………………………..……………adresinde gerçekleşecektir.
    Taraflardan birinin geçerli bir mazeret göstermeksizin ilk toplantıya katılmaması sebebiyle arabuluculuk faaliyetinin sona ermesi durumunda toplantıya katıllmayan taraf, son tutanakta belirtilir ve bu taraf davada kısmen veya tamamen haklı çıksa bile yargılama giderinin tamamından sorumlu tutulur. Ayrıca bu taraf lehine vekâlet ücretine hükmedilmez. Her iki tarafın da ilk toplantıya katılmaması sebebiyle sona eren arabuluculuk faaliyeti üzerine açılacak davalarda tarafların yaptıkları yargılama giderleri kendi üzerlerinde bırakılır (HUAK m.18A/11).

    Arabuluculuk görüşmelerine, gerçek kişilerin kimlik belgesi, şirket yetkililerinin kimlik belgesi ve imza  sirküleri, avukatların kimlik belgesi ve arabuluculuk görüşmelerine katılma konusunda özel yetki bulunan vekâletname ile toplantıya katılması gerekmektedir.
    Arabuluculuk görüşmelerinde idareyi, üst yönetici tarafından belirlenen iki üye ile hukuk birimi amiri veya onun belirleyeceği bir avukat ya da hukuk müşavirinden oluşan komisyon temsil eder (HUAK m. 15/8). Komisyon kendisini vekil ile temsil ettiremez (HUAK Yönetmeliği m. 18/1). (İdarenin taraf olduğu dava şartı arabuluculuk sürecinde kullanılabilir.)


    """+str(isim)+str(adres)+str(basvuruno)+str(telno)+"""




    """
        
        doc= docx.Document()

        doc.add_paragraph(html_str)
        doc.save('Sonuc.docx')
        self.popupmsg("Başarıyla Word Haline Getirildi")




if __name__ == '__main__':
    Arabirim()